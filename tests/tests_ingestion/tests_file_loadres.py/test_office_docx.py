# -*- encoding: utf-8 -*-
"""
@File    :  test_docx.py
@Time    :  2025/08/16 17:20:00
@Author  :  Kevin Wang
@Desc    :  Tests for DoclingDocxLoader (returns LoaderResult)
"""

from datetime import datetime, timezone
from pathlib import Path
from typing import List

import pytest
from docx import Document as PyDocx
from docling_core.types.doc.document import DoclingDocument

from ingestion.base import LoaderResult
from ingestion.file_loaders.docx import DoclingDocxLoader
from objects import FileType


@pytest.fixture
def test_docx_path(tmp_path: Path) -> Path:  # pytest 內建 fixture，會提供臨時目錄並於測試後自動清理
    """建立一個臨時的 DOCX 檔案，供單一檔案測試使用"""
    file_path = tmp_path / "test01.docx"

    doc = PyDocx()
    doc.add_heading("測試文件", level=1)
    doc.add_paragraph("這是一個臨時建立的 DOCX 檔案。")

    props = doc.core_properties
    props.title = "單檔測試標題"
    props.author = "Kevin Wang"
    props.subject = "單檔測試主題"
    now = datetime.now(timezone.utc)
    props.created = now
    props.modified = now

    doc.save(str(file_path))
    return file_path


@pytest.fixture
def test_docx_paths(tmp_path: Path) -> List[Path]:
    """建立多個臨時的 DOCX 檔案（pytest 測試結束後會自動刪除）"""
    paths: List[Path] = []
    now = datetime.now(timezone.utc)

    # file1
    p1 = tmp_path / "test01.docx"
    d1 = PyDocx()
    d1.add_heading("測試文件 01", level=1)
    d1.add_paragraph("內容一。")
    d1.core_properties.title = "批次檔案一"
    d1.core_properties.author = "Kevin Wang"
    d1.core_properties.subject = "批次測試"
    d1.core_properties.created = now
    d1.core_properties.modified = now
    d1.save(str(p1))
    paths.append(p1)

    # file2
    p2 = tmp_path / "test02.docx"
    d2 = PyDocx()
    d2.add_heading("測試文件 02", level=1)
    d2.add_paragraph("內容二。")
    d2.core_properties.title = "批次檔案二"
    d2.core_properties.author = "Kevin Wang"
    d2.core_properties.subject = "批次測試"
    d2.core_properties.created = now
    d2.core_properties.modified = now
    d2.save(str(p2))
    paths.append(p2)

    return paths


class TestDoclingDocxLoader:
    @pytest.fixture
    def docx_loader(self):
        return DoclingDocxLoader()

    def test_load(self, docx_loader, test_docx_path: Path):
        results = docx_loader.load(test_docx_path)
        assert isinstance(results, list) and len(results) == 1

        res = results[0]
        assert isinstance(res, LoaderResult)

        # metadata 檢查（Docx 版本使用 core_properties）
        assert res.metadata.file_type == FileType.DOCX
        assert res.metadata.file_name == test_docx_path.name
        # created_at / modified_at 由 python-docx core_properties 提供
        assert isinstance(res.metadata.created_at, datetime)
        assert isinstance(res.metadata.modified_at, datetime)

        # content 與 doc 基本檢查（包含原文關鍵字）
        assert isinstance(res.content, str) and "測試文件" in res.content
        assert isinstance(res.doc, DoclingDocument)

    def test_load_multi(self, docx_loader, test_docx_paths: List[Path]):
        # 依 markdown 測試邏輯，假設 loader 有提供 load_multi(paths) -> List[LoaderResult]
        results = docx_loader.load_multi(test_docx_paths)

        assert isinstance(results, list)
        assert len(results) == len(test_docx_paths)

        for res, src_path in zip(results, test_docx_paths):
            assert isinstance(res, LoaderResult)

            # metadata 檢查
            assert res.metadata.file_type == FileType.DOCX
            assert res.metadata.file_name == src_path.name
            assert isinstance(res.metadata.created_at, datetime)
            assert isinstance(res.metadata.modified_at, datetime)

            # content 與 doc 基本檢查
            assert isinstance(res.content, str) and "測試" in res.content
            assert isinstance(res.doc, DoclingDocument)

    # TODO: def test_load_real(self, docx_loader):